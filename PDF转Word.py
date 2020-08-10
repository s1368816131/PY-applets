import ctypes,sys
import os
from io import StringIO
from io import open
from concurrent.futures import ProcessPoolExecutor
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from docx import Document

pdf_folder = r'F:\上海wison\PDF'
word_folder = r'F:\上海wison\Word'

def read_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        resource_manager = PDFResourceManager()
        return_str = StringIO()
        lap_params = LAParams()

        device = TextConverter(
            resource_manager,
            return_str,
            laparams=lap_params)
        process_pdf(resource_manager, device, file)
        device.close()

        content = return_str.getvalue()
        return_str.close()
        return content

def process_pdf(rsrcmgr, device, fp, pagenos=None, maxpages=0, password='',
                caching=True, check_extractable=True):
    # Create a PDF parser object associated with the file object.
    parser = PDFParser(fp)
    # Create a PDF document object that stores the document structure.
    doc = PDFDocument(caching=caching)
    # Connect the parser and document objects.
    parser.set_document(doc)
    doc.set_parser(parser)
    # Supply the document password for initialization.
    # (If no password is set, give an empty string.)
    doc.initialize(password)
    # Check if the document allows text extraction. If not, abort.
    if check_extractable and not doc.is_extractable:
        raise PDFTextExtractionNotAllowed('Text extraction is not allowed: %r' % fp)
# Create a PDF interpreter object.
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    # Process each page contained in the document.
    for (pageno,page) in enumerate(doc.get_pages()):
        if pagenos and (pageno not in pagenos): continue
        interpreter.process_page(page)
        if maxpages and maxpages <= pageno+1: break

def save_text_to_word(content, file_path):
    doc = Document()
    for line in content.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)
def pdf_to_word(pdf_file_path, word_file_path):
    content = read_from_pdf(pdf_file_path)
    save_text_to_word(content, word_file_path)
tasks = []
with ProcessPoolExecutor(max_workers=5) as executor:
    for file in os.listdir(pdf_folder):
        extension_name = os.path.splitext(file)[1]
        if extension_name != '.pdf':
            continue
        file_name = os.path.splitext(file)[0]
        pdf_file = pdf_folder + '/' + file
        word_file = word_folder + '/' + file_name + '.docx'
        print('正在处理: ', file)
        result = executor.submit(pdf_to_word, pdf_file, word_file)
        tasks.append(result)
while True:
    exit_flag = True
    for task in tasks:
        if not task.done():
            exit_flag = False
    if exit_flag:
        print('完成')
        exit(0)